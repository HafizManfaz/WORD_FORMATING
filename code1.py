import docx
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
# Open the Word document
doc = docx.Document()

# Get the header of the first section
header = doc.sections[0].header

# Add a table with one row and three columns to the header
table = header.add_table(rows=1, cols=3, width=Inches(6.0))

table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Resize the table to fit the header
for row in table.rows:
    for cell in row.cells:
        cell.width = docx.shared.Inches(2)

# Add bold text to the first cell
first_cell = table.cell(0, 0)
first_cell.text = "Header Data"
first_cell.paragraphs[0].style.font.bold = True
first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


second_cell = table.cell(0, 1)
second_cell.text = "Header Data"
second_cell.alignment = WD_TABLE_ALIGNMENT.CENTER
second_cell.paragraphs[0].style.font.bold = True
second_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



# Add the picture to the third cell
last_cell = table.cell(0, 2)


run = last_cell.paragraphs[0].add_run()
picture = run.add_picture("pic.jpg")
picture.width = docx.shared.Inches(1)
picture.height = docx.shared.Inches(0.5)


# Save the document
doc.save("example.docx")
